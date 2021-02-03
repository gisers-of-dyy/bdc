# -*- encoding:utf-8 -*-
# author:black_liu
# date:2019/02/13
from docx import Document
from docxcompose.composer import Composer
import win32com.client as win32
import win32com.client
from win32com import client
import os
import copy


#获取指定目录下的文件名称
root = "F:\DiYanYuan\python\word\第一次\大于400\\all"
nameList = []
for (dirpath, dirnames, filenames) in os.walk(root):
    for filename in filenames:
        nameList += [os.path.join(dirpath, filename)]
# print(nameList)
nameList.sort()

#生成150个一个文档的list
docxList = []
tempList = []
num = 30
index = 0
while index <len(nameList):
    tempList.append(nameList[index])
    if (index+1) % num == 0 or (index + 1) == len(nameList):
        list1 = copy.deepcopy(tempList)
        docxList.append(list1)
        tempList.clear()
    index += 1

a = 0
while a <len(docxList):
    docName = "F:\DiYanYuan\python\word\第一次\大于400\\all_合并\\" + str(a) + ".docx"
    # 新建空白文档
    new_document = Document()
    composer = Composer(new_document)

    for file in docxList[a]:
        composer.append(Document(file))
    composer.save(docName)
    a += 1


#不正确 重叠了
# a = 0
# while a <len(docxList):
#     docName = "F:\DiYanYuan\python\word\第一次\大于400\\all_合并\\" + str(a) + ".docx"
#     # 启动word应用程序
#     word = win32com.client.Dispatch('Word.Application')
#     word.Visible = False
#     # 新建空白文档
#     new_document = word.Documents.Add()
#     for file in docxList[a]:
#         new_document.Application.Selection.Range.InsertFile(file)
#         # word = Document()
#         # # 添加分页
#         # word.add_page_break()
#     # 保存最终文件，关闭word应用程序
#     new_document.SaveAs(docName)
#     new_document.Close()
#     word.Quit()
#     a += 1

